"""Default Word template for exam notes."""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path


# Color scheme for professional exam notes
COLORS = {
    'primary': RGBColor(0x1A, 0x23, 0x7E),      # Deep indigo
    'secondary': RGBColor(0x30, 0x4F, 0xFE),    # Bright blue
    'accent': RGBColor(0x00, 0xB0, 0xFF),       # Light blue
    'success': RGBColor(0x00, 0x89, 0x7B),      # Teal green
    'warning': RGBColor(0xFF, 0x6D, 0x00),      # Orange
    'danger': RGBColor(0xD3, 0x2F, 0x2F),       # Red
    'text': RGBColor(0x21, 0x21, 0x21),         # Dark gray
    'text_light': RGBColor(0x75, 0x75, 0x75),   # Medium gray
    'bg_light': RGBColor(0xF5, 0xF7, 0xFA),     # Light background
    'concept_bg': RGBColor(0xE8, 0xEA, 0xF6),   # Light indigo background
    'formula_bg': RGBColor(0xFD, 0xF4, 0xE3),   # Light yellow background
    'warning_bg': RGBColor(0xFF, 0xEB, 0xEE),   # Light red background
}


def create_default_template(output_path: Path = None) -> Path:
    """Create a beautiful default template for exam notes."""
    doc = Document()

    # ========== Page Setup ==========
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # ========== Define Styles ==========

    # Title Style (for cover page)
    title_style = doc.styles.add_style('CoverTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = 'Microsoft YaHei'
    title_style.font.size = Pt(28)
    title_style.font.bold = True
    title_style.font.color.rgb = COLORS['primary']
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(12)
    title_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # Subtitle Style
    subtitle_style = doc.styles.add_style('CoverSubtitle', WD_STYLE_TYPE.PARAGRAPH)
    subtitle_style.font.name = 'Microsoft YaHei'
    subtitle_style.font.size = Pt(14)
    subtitle_style.font.color.rgb = COLORS['text_light']
    subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_style.paragraph_format.space_after = Pt(6)
    subtitle_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # Heading 1 - Main sections (核心概念, 公式定理, etc.)
    h1_style = doc.styles['Heading 1']
    h1_style.font.name = 'Microsoft YaHei'
    h1_style.font.size = Pt(16)
    h1_style.font.bold = True
    h1_style.font.color.rgb = COLORS['primary']
    h1_style.paragraph_format.space_before = Pt(18)
    h1_style.paragraph_format.space_after = Pt(10)
    h1_style.paragraph_format.border_bottom.color.rgb = COLORS['accent']
    h1_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # Heading 2 - Subsections
    h2_style = doc.styles['Heading 2']
    h2_style.font.name = 'Microsoft YaHei'
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True
    h2_style.font.color.rgb = COLORS['secondary']
    h2_style.paragraph_format.space_before = Pt(14)
    h2_style.paragraph_format.space_after = Pt(8)
    h2_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # Heading 3 - Sub-subsections
    h3_style = doc.styles['Heading 3']
    h3_style.font.name = 'Microsoft YaHei'
    h3_style.font.size = Pt(12)
    h3_style.font.bold = True
    h3_style.font.color.rgb = COLORS['accent']
    h3_style.paragraph_format.space_before = Pt(10)
    h3_style.paragraph_format.space_after = Pt(6)
    h3_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # Normal - Body text
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Microsoft YaHei'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = COLORS['text']
    normal_style.paragraph_format.line_spacing = 1.6
    normal_style.paragraph_format.space_after = Pt(6)
    normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # List Bullet
    bullet_style = doc.styles['List Bullet']
    bullet_style.font.name = 'Microsoft YaHei'
    bullet_style.font.size = Pt(11)
    bullet_style.font.color.rgb = COLORS['text']
    bullet_style.paragraph_format.left_indent = Cm(0.75)
    bullet_style.paragraph_format.space_after = Pt(4)
    bullet_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # List Number
    number_style = doc.styles['List Number']
    number_style.font.name = 'Microsoft YaHei'
    number_style.font.size = Pt(11)
    number_style.font.color.rgb = COLORS['text']
    number_style.paragraph_format.left_indent = Cm(0.75)
    number_style.paragraph_format.space_after = Pt(4)
    number_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # Concept Style - for core concepts
    concept_style = doc.styles.add_style('Concept', WD_STYLE_TYPE.PARAGRAPH)
    concept_style.font.name = 'Microsoft YaHei'
    concept_style.font.size = Pt(11)
    concept_style.font.bold = True
    concept_style.font.color.rgb = COLORS['primary']
    concept_style.paragraph_format.left_indent = Cm(0.5)
    concept_style.paragraph_format.space_before = Pt(6)
    concept_style.paragraph_format.space_after = Pt(6)
    concept_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # Formula Style - for mathematical formulas
    formula_style = doc.styles.add_style('Formula', WD_STYLE_TYPE.PARAGRAPH)
    formula_style.font.name = 'Cambria Math'
    formula_style.font.size = Pt(12)
    formula_style.font.italic = True
    formula_style.font.color.rgb = COLORS['text']
    formula_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formula_style.paragraph_format.left_indent = Cm(1)
    formula_style.paragraph_format.right_indent = Cm(1)
    formula_style.paragraph_format.space_before = Pt(8)
    formula_style.paragraph_format.space_after = Pt(8)

    # Warning Style - for important warnings
    warning_style = doc.styles.add_style('Warning', WD_STYLE_TYPE.PARAGRAPH)
    warning_style.font.name = 'Microsoft YaHei'
    warning_style.font.size = Pt(11)
    warning_style.font.bold = True
    warning_style.font.color.rgb = COLORS['danger']
    warning_style.paragraph_format.left_indent = Cm(0.5)
    warning_style.paragraph_format.space_before = Pt(6)
    warning_style.paragraph_format.space_after = Pt(6)
    warning_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # ExamPoint Style - for exam points
    exam_style = doc.styles.add_style('ExamPoint', WD_STYLE_TYPE.PARAGRAPH)
    exam_style.font.name = 'Microsoft YaHei'
    exam_style.font.size = Pt(11)
    exam_style.font.bold = True
    exam_style.font.color.rgb = COLORS['success']
    exam_style.paragraph_format.left_indent = Cm(0.5)
    exam_style.paragraph_format.space_before = Pt(4)
    exam_style.paragraph_format.space_after = Pt(4)
    exam_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # Quote Style - for memorable tips
    quote_style = doc.styles.add_style('Quote', WD_STYLE_TYPE.PARAGRAPH)
    quote_style.font.name = 'Microsoft YaHei'
    quote_style.font.size = Pt(11)
    quote_style.font.italic = True
    quote_style.font.color.rgb = COLORS['text_light']
    quote_style.paragraph_format.left_indent = Cm(1)
    quote_style.paragraph_format.right_indent = Cm(1)
    quote_style.paragraph_format.space_before = Pt(8)
    quote_style.paragraph_format.space_after = Pt(8)
    quote_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # ========== Setup Header/Footer ==========
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = "备考专用笔记"
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_para.style.font.size = Pt(9)
    header_para.style.font.color.rgb = COLORS['text_light']

    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = ""
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.style.font.size = Pt(9)
    footer_para.style.font.color.rgb = COLORS['text_light']

    # Save template
    if output_path:
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        return output_path

    return doc


def get_default_template_path() -> Path:
    """Get the path to the default template file."""
    from pathlib import Path
    template_dir = Path(__file__).parent.parent.parent / "config" / "templates"
    template_dir.mkdir(parents=True, exist_ok=True)
    return template_dir / "default_template.docx"


def ensure_default_template() -> Path:
    """Ensure default template exists and return its path."""
    template_path = get_default_template_path()
    if not template_path.exists():
        create_default_template(template_path)
    return template_path


if __name__ == "__main__":
    # Create template in config/templates directory
    template_path = get_default_template_path()
    create_default_template(template_path)
    print(f"Template created: {template_path}")
