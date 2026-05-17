"""Word document generation module for exam preparation notes."""

import logging
import re
from datetime import datetime
from pathlib import Path
from typing import Optional, List

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

from pdf_summarizer.latex_to_omml import latex_to_omml, latex_to_omath_para
from pdf_summarizer.models import SummaryOutput, SummarySection

logger = logging.getLogger(__name__)


class NoteStyles:
    """Predefined styles for lecture notes."""

    # Colors
    TITLE_COLOR = RGBColor(0x00, 0x00, 0x00)       # Black
    HEADING1_COLOR = RGBColor(0x33, 0x33, 0x33)    # Dark gray
    HEADING2_COLOR = RGBColor(0x55, 0x55, 0x55)    # Medium gray
    ACCENT_COLOR = RGBColor(0x33, 0x33, 0x33)      # Dark gray
    SUCCESS_COLOR = RGBColor(0x2E, 0x7D, 0x32)     # Green
    WARNING_COLOR = RGBColor(0xCC, 0x66, 0x00)     # Brown
    CONCEPT_BG = RGBColor(0xF5, 0xF5, 0xF5)        # Light gray background
    FORMULA_BG = RGBColor(0xFA, 0xFA, 0xFA)        # Very light gray background

    # Font sizes
    TITLE_SIZE = Pt(22)
    HEADING1_SIZE = Pt(16)
    HEADING2_SIZE = Pt(14)
    HEADING3_SIZE = Pt(12)
    BODY_SIZE = Pt(11)
    SMALL_SIZE = Pt(9)

    # Fonts
    CHINESE_FONT = "Microsoft YaHei"
    MONO_FONT = "Consolas"


class NoteGenerator:
    """
    Generate professional Word documents for lecture notes.

    Features:
    - Customizable templates
    - Styled content blocks (definitions, theorems, formulas)
    - Automatic markdown parsing
    - Cover page generation
    - Header/footer with metadata
    """

    def __init__(self, template_path: Optional[Path] = None):
        """
        Initialize the note generator.

        Args:
            template_path: Optional path to a .docx template file
        """
        self.template_path = template_path
        self.doc: Optional[Document] = None
        self.styles = NoteStyles()
        self._metadata: dict = {}

        # Import and use default template if none specified
        if template_path is None:
            try:
                from pdf_summarizer.default_template import ensure_default_template
                self.template_path = ensure_default_template()
            except Exception:
                pass  # Fall back to default styling

    def _create_document(self) -> Document:
        """Create a new document, optionally from template."""
        if self.template_path and self.template_path.exists():
            doc = Document(str(self.template_path))
        else:
            doc = Document()

        # Setup default styles
        self._setup_styles(doc)
        return doc

    def _setup_styles(self, doc: Document):
        """Configure document styles for Chinese text."""
        # Set default font for Normal style
        style = doc.styles['Normal']
        style.font.name = self.styles.CHINESE_FONT
        style.font.size = self.styles.BODY_SIZE
        style._element.rPr.rFonts.set(qn('w:eastAsia'), self.styles.CHINESE_FONT)

        # Configure heading styles
        for i, (size, color) in enumerate([
            (self.styles.HEADING1_SIZE, self.styles.HEADING1_COLOR),
            (self.styles.HEADING2_SIZE, self.styles.HEADING2_COLOR),
            (self.styles.HEADING3_SIZE, self.styles.HEADING2_COLOR),
        ], start=1):
            style = doc.styles[f'Heading {i}']
            style.font.name = self.styles.CHINESE_FONT
            style.font.size = size
            style.font.color.rgb = color
            style.font.bold = True
            style._element.rPr.rFonts.set(qn('w:eastAsia'), self.styles.CHINESE_FONT)

    def add_cover(self, title: str, metadata: Optional[dict] = None):
        """
        Add a cover page to the document.

        Args:
            title: Main title for the cover
            metadata: Optional metadata (source, date, etc.)
        """
        if self.doc is None:
            self.doc = self._create_document()

        metadata = metadata or {}
        self._metadata = metadata

        # Add title
        title_para = self.doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title)
        title_run.font.size = self.styles.TITLE_SIZE
        title_run.font.color.rgb = self.styles.TITLE_COLOR
        title_run.font.bold = True

        # Add subtitle
        if metadata.get('subtitle'):
            subtitle = self.doc.add_paragraph()
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub_run = subtitle.add_run(metadata['subtitle'])
            sub_run.font.size = Pt(14)
            sub_run.font.color.rgb = self.styles.HEADING2_COLOR

        # Add separator
        self.doc.add_paragraph()
        self._add_separator()
        self.doc.add_paragraph()

        # Add metadata table
        if metadata:
            table = self.doc.add_table(rows=0, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            meta_items = {
                '来源文件': metadata.get('source', '-'),
                '生成时间': datetime.now().strftime('%Y-%m-%d %H:%M'),
                'AI 模型': metadata.get('model', '-'),
                '页数': metadata.get('pages', '-'),
            }

            for key, value in meta_items.items():
                row = table.add_row()
                row.cells[0].text = key
                row.cells[1].text = str(value)
                row.cells[0].paragraphs[0].runs[0].font.bold = True

        # Page break after cover
        self.doc.add_page_break()

    def add_summary(self, summary_md: str):
        """
        Parse markdown content and render as Word document.

        Args:
            summary_md: Markdown formatted content
        """
        if self.doc is None:
            self.doc = self._create_document()

        # Parse and render content
        self._render_markdown(summary_md)

    def _render_markdown(self, content: str):
        """Parse markdown and render to Word."""
        lines = content.split('\n')
        i = 0

        while i < len(lines):
            line = lines[i].strip()

            # Empty line
            if not line:
                i += 1
                continue

            # Section marker 【xxx】
            section_match = re.match(r'【([^】]+)】', line)
            if section_match:
                section_title = section_match.group(1)
                self._add_section_header(section_title)
                # Add remaining content on the line
                remaining = line[section_match.end():].strip()
                if remaining:
                    self._render_line(remaining)
                i += 1
                continue

            # Markdown heading
            heading_match = re.match(r'^(#{1,3})\s+(.+)$', line)
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2)
                self.doc.add_heading(text, level=level)
                i += 1
                continue

            # Regular content
            self._render_line(line)
            i += 1

    def _render_line(self, line: str):
        """Render a single line based on content type."""
        # Check for different content types
        if line.startswith(('-', '•', '*', '·')):
            self._add_bullet_item(line)
        elif re.match(r'^\d+[\.、\)]\s', line):
            self._add_numbered_item(line)
        elif self._is_formula(line):
            self._add_formula_block(line)
        else:
            self._add_paragraph(line)

    def _add_section_header(self, title: str):
        """Add a styled section header."""
        heading = self.doc.add_heading(title, level=1)

    def _add_concept_box(self, concepts: List[str]):
        """Add a highlighted box for core concepts."""
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        cell = table.cell(0, 0)
        cell.text = ""

        # Set background color
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'E3F2FD')
        cell._tc.get_or_add_tcPr().append(shading)

        # Add content
        for concept in concepts:
            para = cell.add_paragraph()
            run = para.add_run(f"• {concept}")
            run.font.size = Pt(11)

        self.doc.add_paragraph()

    def _add_formula_block(self, text: str):
        """Add a formatted formula/theorem block with a real Word equation."""
        # Strip outer markers
        formula = text
        for prefix, suffix in [('$$', '$$'), ('$', '$'), ('\\[', '\\]'), ('\\(', '\\)')]:
            if formula.startswith(prefix) and formula.endswith(suffix):
                formula = formula[len(prefix):-len(suffix)].strip()
                break

        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add light background effect
        para.paragraph_format.left_indent = Cm(1)
        para.paragraph_format.right_indent = Cm(1)
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after = Pt(8)

        try:
            omml_para = latex_to_omath_para(formula)
            para._element.append(omml_para)
        except Exception:
            # Fallback: use _fallback_omath for safe XML handling
            from pdf_summarizer.latex_to_omml import _fallback_omath
            omml = _fallback_omath(formula)
            para._element.append(omml)

    def _render_latex_formula(self, para, latex: str):
        """Render LaTeX formula as a real Word equation."""
        latex = latex.strip()
        if latex.startswith('$$') and latex.endswith('$$'):
            latex = latex[2:-2].strip()
        elif latex.startswith('$') and latex.endswith('$'):
            latex = latex[1:-1].strip()

        omml = latex_to_omml(latex)
        para._element.append(omml)

    def _render_formula_text(self, para, text: str):
        """Render formula text as a real Word equation, fallback to monospace."""
        # Strip outer markers if present
        formula = text
        for prefix, suffix in [('$$', '$$'), ('$', '$'), ('\\[', '\\]'), ('\\(', '\\)')]:
            if formula.startswith(prefix) and formula.endswith(suffix):
                formula = formula[len(prefix):-len(suffix)].strip()
                break

        if formula:
            try:
                omml_para = latex_to_omath_para(formula)
                para._element.append(omml_para)
                return
            except Exception:
                pass

        # Fallback: use _fallback_omath for safe XML handling
        from pdf_summarizer.latex_to_omml import _fallback_omath
        omml = _fallback_omath(text)
        para._element.append(omml)

    def _latex_to_word_text(self, latex: str) -> str:
        """Convert LaTeX syntax to readable text for Word."""
        # Common LaTeX conversions
        conversions = {
            r'\mathbf': '',
            r'\mathbf{': '',
            r'\psi': 'ψ',
            r'\Psi': 'Ψ',
            r'\phi': 'φ',
            r'\Phi': 'Φ',
            r'\theta': 'θ',
            r'\Theta': 'Θ',
            r'\omega': 'ω',
            r'\Omega': 'Ω',
            r'\alpha': 'α',
            r'\beta': 'β',
            r'\gamma': 'γ',
            r'\Gamma': 'Γ',
            r'\delta': 'δ',
            r'\Delta': 'Δ',
            r'\epsilon': 'ε',
            r'\lambda': 'λ',
            r'\Lambda': 'Λ',
            r'\mu': 'μ',
            r'\nu': 'ν',
            r'\pi': 'π',
            r'\Pi': 'Π',
            r'\sigma': 'σ',
            r'\Sigma': 'Σ',
            r'\tau': 'τ',
            r'\eta': 'η',
            r'\rho': 'ρ',
            r'\times': '×',
            r'\cdot': '·',
            r'\div': '÷',
            r'\pm': '±',
            r'\infty': '∞',
            r'\int': '∫',
            r'\sum': '∑',
            r'\prod': '∏',
            r'\sqrt': '√',
            r'\approx': '≈',
            r'\neq': '≠',
            r'\leq': '≤',
            r'\geq': '≥',
            r'\left': '',
            r'\right': '',
            r'\{': '{',
            r'\}': '}',
            r'\frac': '/',
        }

        result = latex
        for latex_cmd, replacement in conversions.items():
            result = result.replace(latex_cmd, replacement)

        # Handle fractions: \frac{a}{b} -> a/b
        result = re.sub(r'\{([^}]*)\}\{([^}]*)\}', r'\1/\2', result)

        # Handle simple braces removal
        result = re.sub(r'\{([^}]*)\}', r'\1', result)

        return result.strip()

    def _latex_command_to_symbol(self, cmd: str) -> str:
        """Convert a LaTeX command to its symbol."""
        cmd = cmd.rstrip('{}')

        symbol_map = {
            '\\alpha': 'α', '\\beta': 'β', '\\gamma': 'γ', '\\delta': 'δ',
            '\\epsilon': 'ε', '\\zeta': 'ζ', '\\eta': 'η', '\\theta': 'θ',
            '\\iota': 'ι', '\\kappa': 'κ', '\\lambda': 'λ', '\\mu': 'μ',
            '\\nu': 'ν', '\\xi': 'ξ', '\\pi': 'π', '\\rho': 'ρ',
            '\\sigma': 'σ', '\\tau': 'τ', '\\upsilon': 'υ', '\\phi': 'φ',
            '\\chi': 'χ', '\\psi': 'ψ', '\\omega': 'ω',
            '\\Gamma': 'Γ', '\\Delta': 'Δ', '\\Theta': 'Θ', '\\Lambda': 'Λ',
            '\\Xi': 'Ξ', '\\Pi': 'Π', '\\Sigma': 'Σ', '\\Phi': 'Φ',
            '\\Psi': 'Ψ', '\\Omega': 'Ω',
            '\\times': '×', '\\cdot': '·', '\\div': '÷',
            '\\pm': '±', '\\mp': '∓', '\\infty': '∞',
            '\\int': '∫', '\\sum': '∑', '\\prod': '∏',
            '\\sqrt': '√', '\\approx': '≈', '\\neq': '≠',
            '\\leq': '≤', '\\geq': '≥', '\\ll': '≪', '\\gg': '≫',
            '\\leftarrow': '←', '\\rightarrow': '→', '\\Leftarrow': '⇐', '\\Rightarrow': '⇒',
            '\\leftrightarrow': '↔', '\\Leftrightarrow': '⇔',
            '\\partial': '∂', '\\nabla': '∇', '\\infty': '∞',
            '\\forall': '∀', '\\exists': '∃', '\\in': '∈', '\\notin': '∉',
            '\\subset': '⊂', '\\supset': '⊃', '\\subseteq': '⊆', '\\supseteq': '⊇',
            '\\cup': '∪', '\\cap': '∩', '\\emptyset': '∅',
            '\\mathbb{R}': 'ℝ', '\\mathbb{Z}': 'ℤ', '\\mathbb{N}': 'ℕ',
            '\\mathbb{Q}': 'ℚ', '\\mathbb{C}': 'ℂ',
            '\\vec': '→', '\\hat': '^', '\\bar': '¯',
            '\\cdot': '·', '\\circ': '∘', '\\star': '★',
        }

        return symbol_map.get(cmd, cmd)

    def _add_warning_box(self, text: str):
        """Add a warning/highlight box."""
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        cell = table.cell(0, 0)
        para = cell.paragraphs[0]
        run = para.add_run(f"⚠️ {text}")
        run.font.color.rgb = self.styles.WARNING_COLOR
        run.font.bold = True

        # Warning background
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'FFF3E0')
        cell._tc.get_or_add_tcPr().append(shading)

        self.doc.add_paragraph()

    def _add_exam_point(self, text: str, number: int):
        """Add a highlighted exam point."""
        para = self.doc.add_paragraph(style='List Number')

        run = para.add_run(text)
        run.font.color.rgb = self.styles.SUCCESS_COLOR
        run.font.bold = True

    def _add_bullet_item(self, text: str):
        """Add a bullet list item with consistent formatting."""
        # Remove bullet marker
        cleaned = re.sub(r'^[-•*·]\s*', '', text)

        para = self.doc.add_paragraph(style='List Bullet')
        para.paragraph_format.line_spacing = 1.5

        # Check for LaTeX content (with $ delimiters or raw LaTeX commands)
        has_latex = ('$' in cleaned or '\\(' in cleaned or '\\[' in cleaned
                     or re.search(r'\\(?:frac|int|sum|prod|sqrt|mathbf|vec|hat|bar|psi|phi|theta|omega|alpha|beta|gamma|delta|epsilon|lambda|mu|pi|sigma|tau|hbar|partial|nabla|infty|times|cdot|div|pm|approx|neq|leq|geq)', cleaned))

        if has_latex:
            self._add_text_with_latex(para, cleaned)
        else:
            # Check if this is a special section
            if hasattr(self, '_current_section_type'):
                if self._current_section_type == 'warning':
                    run = para.add_run(cleaned)
                    run.font.color.rgb = self.styles.WARNING_COLOR
                    run.font.size = self.styles.BODY_SIZE
                elif self._current_section_type == 'exam':
                    run = para.add_run(cleaned)
                    run.font.color.rgb = self.styles.SUCCESS_COLOR
                    run.font.size = self.styles.BODY_SIZE
                else:
                    run = para.add_run(cleaned)
                    run.font.size = self.styles.BODY_SIZE
            else:
                run = para.add_run(cleaned)
                run.font.size = self.styles.BODY_SIZE

            # Set East Asian font
            for run in para.runs:
                run.font.name = self.styles.CHINESE_FONT
                run._element.rPr.rFonts.set(qn('w:eastAsia'), self.styles.CHINESE_FONT)

    def _add_numbered_item(self, text: str):
        """Add a numbered list item with consistent formatting."""
        para = self.doc.add_paragraph(style='List Number')
        para.paragraph_format.line_spacing = 1.5
        cleaned = re.sub(r'^\d+[\.、\)]\s*', '', text)

        # Check for LaTeX in numbered items
        has_latex = ('$' in cleaned or '\\(' in cleaned or '\\[' in cleaned
                     or re.search(r'\\(?:frac|int|sum|prod|sqrt|mathbf|vec|hat|bar|psi|phi|theta|omega|alpha|beta|gamma|delta|epsilon|lambda|mu|pi|sigma|tau|hbar|partial|nabla|infty|times|cdot|div|pm|approx|neq|leq|geq)', cleaned))
        if has_latex:
            self._add_text_with_latex(para, cleaned)
        else:
            run = para.add_run(cleaned)
            run.font.name = self.styles.CHINESE_FONT
            run.font.size = self.styles.BODY_SIZE
            run._element.rPr.rFonts.set(qn('w:eastAsia'), self.styles.CHINESE_FONT)

    def _add_text_with_latex(self, para, text: str):
        """Add text with inline LaTeX formulas rendered as real Word equations."""
        # Split by LaTeX delimiters
        parts = re.split(r'(\$\$.*?\$\$|\$[^\$]+\$|\\\[.*?\\\]|\\\(.*?\\\))', text)

        has_delimiters = any(
            (p.startswith('$$') and p.endswith('$$')) or
            (p.startswith('$') and p.endswith('$') and len(p) > 1) or
            (p.startswith('\\[') and p.endswith('\\]')) or
            (p.startswith('\\(') and p.endswith('\\)'))
            for p in parts
        )

        # No delimiters but raw LaTeX → try full-text equation, fallback monospace
        if not has_delimiters and re.search(r'\\(?:frac|int|sum|prod|sqrt|mathbf|vec|hat|bar|psi|phi|theta|omega|alpha|beta|gamma|delta|epsilon|lambda|mu|pi|sigma|tau|hbar|partial|nabla|infty|times|cdot|div|pm|approx|neq|leq|geq|left|right)', text):
            try:
                omml = latex_to_omml(text)
                para._element.append(omml)
            except Exception:
                run = para.add_run(text)
                run.font.name = self.styles.MONO_FONT
                run.font.size = Pt(10)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), self.styles.CHINESE_FONT)
            return

        for part in parts:
            if not part:
                continue

            # Display math $$...$$ or inline $...$
            if (part.startswith('$$') and part.endswith('$$')) or \
               (part.startswith('$') and part.endswith('$') and len(part) > 1):
                formula = part.strip('$')
                try:
                    omml = latex_to_omml(formula)
                    para._element.append(omml)
                except Exception:
                    run = para.add_run(part)
                    run.font.name = self.styles.MONO_FONT
                    run.font.size = Pt(10)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), self.styles.CHINESE_FONT)
            # LaTeX display \[...\] or inline \(...\)
            elif (part.startswith('\\[') and part.endswith('\\]')) or \
                 (part.startswith('\\(') and part.endswith('\\)')):
                formula = part[2:-2].strip()
                try:
                    omml = latex_to_omml(formula)
                    para._element.append(omml)
                except Exception:
                    run = para.add_run(part)
                    run.font.name = self.styles.MONO_FONT
                    run.font.size = Pt(10)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), self.styles.CHINESE_FONT)
            # Regular text
            else:
                run = para.add_run(part)
                run.font.name = self.styles.CHINESE_FONT
                run.font.size = self.styles.BODY_SIZE
                run._element.rPr.rFonts.set(qn('w:eastAsia'), self.styles.CHINESE_FONT)

    def _is_formula(self, text: str) -> bool:
        """Detect if text is a formula."""
        formula_patterns = [
            r'\$\$.*?\$\$',         # Display math $$...$$
            r'\\\[.*?\\\]',         # LaTeX display math \[...\]
            r'\\\(.*?\\\)',         # LaTeX inline math \(...\)
            r'\\[a-zA-Z]+\{',       # LaTeX commands with braces
            r'[=≈≠<>≤≥]',           # Math operators
            r'[∑∫∏√]',              # Math symbols
            r'\b[a-z]_\{',          # Subscripts like a_{ij}
            r'\^[\{\^]',            # Superscripts
            r'[α-ωΑ-Ω]',            # Greek letters
            r'\\mathbf',            # Bold math
            r'\\frac',              # Fractions
            r'\\psi',               # Common math symbols
            r'\\phi',
            r'\\theta',
            r'\\omega',
            r'\\sum',
            r'\\int',
        ]
        return any(re.search(p, text) for p in formula_patterns)

    def _add_paragraph(self, text: str):
        """Add a normal paragraph with inline equations rendered properly."""
        para = self.doc.add_paragraph()
        para.paragraph_format.first_line_indent = Cm(0.5)
        para.paragraph_format.line_spacing = 1.5

        # Check for inline LaTeX
        if '$' in text or '\\(' in text or '\\[' in text:
            # Split by LaTeX delimiters
            parts = re.split(r'(\$\$.*?\$\$|\$[^\$]+\$|\\\[.*?\\\]|\\\(.*?\\\))', text)

            for part in parts:
                if not part:
                    continue

                # Display math $$...$$ or inline $...$
                if (part.startswith('$$') and part.endswith('$$')) or \
                   (part.startswith('$') and part.endswith('$') and len(part) > 1):
                    formula = part.strip('$')
                    try:
                        omml = latex_to_omml(formula)
                        para._element.append(omml)
                    except Exception:
                        run = para.add_run(part)
                        run.font.name = self.styles.MONO_FONT
                        run.font.size = Pt(10)
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), self.styles.CHINESE_FONT)
                # LaTeX display \[...\] or inline \(...\)
                elif (part.startswith('\\[') and part.endswith('\\]')) or \
                     (part.startswith('\\(') and part.endswith('\\)')):
                    formula = part[2:-2].strip()
                    try:
                        omml = latex_to_omml(formula)
                        para._element.append(omml)
                    except Exception:
                        run = para.add_run(part)
                        run.font.name = self.styles.MONO_FONT
                        run.font.size = Pt(10)
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), self.styles.CHINESE_FONT)
                # Regular text
                else:
                    run = para.add_run(part)
                    run.font.name = self.styles.CHINESE_FONT
                    run.font.size = self.styles.BODY_SIZE
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), self.styles.CHINESE_FONT)
        else:
            # Regular text without LaTeX
            run = para.add_run(text)
            run.font.name = self.styles.CHINESE_FONT
            run.font.size = self.styles.BODY_SIZE
            run._element.rPr.rFonts.set(qn('w:eastAsia'), self.styles.CHINESE_FONT)

    def _add_separator(self):
        """Add a horizontal separator line."""
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run('─' * 50)
        run.font.color.rgb = self.styles.HEADING2_COLOR

    def add_header_footer(self, title: str, date: Optional[str] = None):
        """Add header and footer to all sections."""
        if self.doc is None:
            return

        date_str = date or datetime.now().strftime('%Y-%m-%d')

        for section in self.doc.sections:
            # Header
            header = section.header
            header_para = header.paragraphs[0]
            header_para.text = f"{title} | 备考笔记"
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Footer
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.text = f"生成日期: {date_str}"
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def save(self, output_path: Path) -> Path:
        """
        Save the document to file.

        Args:
            output_path: Path to save the document

        Returns:
            Path to the saved file
        """
        if self.doc is None:
            raise ValueError("No document to save. Call add_cover() or add_summary() first.")

        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        # Add header/footer if metadata exists
        if self._metadata.get('title'):
            self.add_header_footer(self._metadata['title'])

        self.doc.save(str(output_path))
        logger.info(f"Document saved: {output_path}")

        return output_path


class DocxWriter(NoteGenerator):
    """Legacy compatible writer."""

    def write(self, summary: SummaryOutput, output_path: Path) -> Path:
        """Write summary to Word document."""
        title = summary.metadata.get('title', summary.source_file) if summary.metadata else summary.source_file

        # Create document
        self.add_cover(
            title=title,
            metadata={
                'source': summary.source_file,
                'subtitle': '',
            }
        )

        # Add content
        if summary.raw_response:
            self.add_summary(summary.raw_response)
        else:
            # Fallback to structured content
            for section in summary.sections:
                self.doc.add_heading(section.title, level=section.level)
                self.add_summary(section.content)

        return self.save(output_path)


def generate_batch(
    notes_data: List[dict],
    output_dir: Path,
    template_path: Optional[Path] = None,
    generate_index: bool = True,
) -> List[Path]:
    """
    Generate multiple Word documents from notes data.

    Args:
        notes_data: List of note data dicts with 'title', 'content', 'metadata'
        output_dir: Directory to save all documents
        template_path: Optional template file
        generate_index: Whether to generate an index document

    Returns:
        List of paths to generated documents
    """
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    generated_files: List[Path] = []
    index_entries: List[dict] = []

    for note_data in notes_data:
        generator = NoteGenerator(template_path=template_path)

        # Add cover
        generator.add_cover(
            title=note_data.get('title', '未命名笔记'),
            metadata=note_data.get('metadata', {})
        )

        # Add content
        generator.add_summary(note_data.get('content', ''))

        # Save
        safe_title = re.sub(r'[\\/:*?"<>|]', '_', note_data.get('title', 'note'))
        output_path = output_dir / f"{safe_title}_笔记.docx"
        generator.save(output_path)

        generated_files.append(output_path)
        index_entries.append({
            'title': note_data.get('title'),
            'path': output_path.name,
            'source': note_data.get('metadata', {}).get('source', '-'),
        })

        logger.info(f"Generated: {output_path}")

    # Generate index document
    if generate_index and index_entries:
        _generate_index_document(index_entries, output_dir)

    return generated_files


def _generate_index_document(entries: List[dict], output_dir: Path):
    """Generate an index document listing all notes."""
    doc = Document()

    # Title
    title = doc.add_heading("笔记索引", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"共计 {len(entries)} 份笔记")
    doc.add_paragraph()

    # Table of contents
    table = doc.add_table(rows=0, cols=3)
    table.style = 'Table Grid'

    # Header row
    header = table.add_row()
    header.cells[0].text = "序号"
    header.cells[1].text = "笔记名称"
    header.cells[2].text = "来源文件"

    for cell in header.cells:
        cell.paragraphs[0].runs[0].font.bold = True

    # Content rows
    for i, entry in enumerate(entries, 1):
        row = table.add_row()
        row.cells[0].text = str(i)
        row.cells[1].text = entry['title']
        row.cells[2].text = entry['source']

    # Save
    index_path = output_dir / "笔记索引.docx"
    doc.save(str(index_path))
    logger.info(f"Index saved: {index_path}")


def write_summary(summary: SummaryOutput, output_path: Path) -> Path:
    """Convenience function to write summary to document."""
    writer = DocxWriter()
    return writer.write(summary, output_path)
